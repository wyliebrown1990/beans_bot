import os
import json
import logging
import glob
from sqlalchemy.orm import sessionmaker
from flask import Flask, current_app
from flask_session import Session
from app.models import JobDescriptions, InterviewHistory, Resumes, Users
from app.database import get_db
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import SystemMessage, HumanMessage
import requests
import fitz  # PyMuPDF for PDF processing
import docx
import re
from datetime import datetime
from sqlalchemy import create_engine, inspect, insert
from sqlalchemy.orm import scoped_session
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Initialize the OpenAI chat model
openai_api_key = os.getenv("OPENAI_API_KEY")
model = ChatOpenAI(api_key=openai_api_key)

def update_process_status(app, user_id, status):
    with app.app_context():
        try:
            logging.debug(f"Updating status for user_id: {user_id} to status: {status}")
            response = requests.post('http://localhost:5011/update_status', json={
                'user_id': user_id,
                'status': status
            })
            logging.debug(f"Status update response: {response.status_code} - {response.text}")
            if response.status_code != 200:
                logging.error(f"Failed to update status for user_id: {user_id}. Status: {status}")
        except Exception as e:
            logging.error(f"Exception while updating status for user_id: {user_id}. Status: {status}. Error: {str(e)}")

def process_text(app, text, user_id):
    with app.app_context():
        try:
            update_process_status(app, user_id, "Processing text submission")
            cleaned_text = text.strip()  # Simple text cleaning
            analysis = get_job_description_analysis(cleaned_text)
            store_job_description_data(analysis, user_id)
            update_process_status(app, user_id, "Processing complete")
        except Exception as e:
            logging.error(f"Error in process_text: {str(e)}")
            update_process_status(app, user_id, f"Processing error: {str(e)}")

def process_file(app, file_path, user_id):
    with app.app_context():
        try:
            update_process_status(app, user_id, "Processing file upload")
            if file_path.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                extracted_text = extract_text_from_docx(file_path)
            elif file_path.endswith('.txt'):
                with open(file_path, 'r') as f:
                    extracted_text = f.read()
            else:
                raise ValueError("Unsupported file type")

            cleaned_text = extracted_text.strip()
            
            # Determine if the file is a job description or a resume
            if "resume" in file_path.lower():
                update_process_status(app, user_id, "Analyzing resume")
                analysis = get_resume_analysis(model, cleaned_text)
                store_resume_analysis(user_id, analysis)
            else:
                update_process_status(app, user_id, "Analyzing job description")
                analysis = get_job_description_analysis(cleaned_text)
                store_job_description_data(analysis, user_id)
                
            update_process_status(app, user_id, "Processing complete")
        except Exception as e:
            logging.error(f"Error in process_file: {str(e)}")
            update_process_status(app, user_id, f"Processing error: {str(e)}")

def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def get_job_description_analysis(job_description_text):
    openai_api_key = os.getenv("OPENAI_API_KEY")
    model = ChatOpenAI(api_key=openai_api_key, model="gpt-3.5-turbo")

    prompt_text = (
    "You are a professional Job Description analyst. Your job is to take the job description that I am sending you in my user message and extract the details below. "
    "Please extract the relevant information from the following job description and return the results in JSON format. If a job description is missing any detail, then return a JSON value of null for that field. "
    "Desired Output in JSON Format: {{\"job_details\": {{\"title\": \"\", \"level\": \"\", \"location\": \"\", \"type\": \"\", \"salary\": \"\", \"responsibilities\": [], \"personal_qualifications\": []}}, "
    "\"company_information\": {{\"name\": \"\", \"size\": \"\", \"industry\": \"\", \"mission_and_values\": \"\"}}, "
    "\"requirements_and_qualifications\": {{\"education_background\": [], \"required_professional_experiences\": [], \"nice_to_have_experiences\": []}}, "
    "\"required_skill_sets\": [], \"required_technical_skills\": [], \"required_soft_skills\": [], \"keywords_analysis\": []}}"
    "\n"
    "1. Job Title: Extract the job title (e.g., Senior Software Engineer). "
    "2. Job Level: Extract the job level (e.g., Mid-Senior level). "
    "3. Job Location: Extract the job location (e.g., New York, NY (Hybrid)). "
    "4. Job Type: Extract the job type (e.g., Full-time). "
    "5. Job Salary: Extract the salary range (e.g., $175K/yr - $205K/yr). "
    "6. Job Responsibilities: List the responsibilities mentioned in the job description. "
    "7. Personal Qualifications: List the personal qualifications mentioned in the job description. "
    "8. Company Name: Extract the company name (e.g., K Health). "
    "9. Company Size: Extract the company size (e.g., 201-500 employees). "
    "10. Company Industry: Extract the industry the company operates in (e.g., Telehealth and AI Healthcare). "
    "11. Company Mission and Values: Summarize the company's mission and values (e.g., Use the power of AI to get everyone access to higher quality healthcare at more affordable costs). "
    "12. Education Background: Summarize the educational background required (e.g., Bachelor's degree in Computer Science, Engineering, or a related field). "
    "13. Required Professional Experiences: List the required professional experiences (e.g., 5+ years of software engineering experience, Experience with highly-scalable, distributed systems). "
    "14. Nice to Have Experiences: List any nice to have experiences (e.g., Experience with modern cloud technologies such as Docker, Kubernetes, Kafka, GCP/AWS suite). "
    "15. Required Skill Sets: List the required skill sets (e.g., Node.js, TypeScript, GraphQL, Apollo Federation, Problem Solving, Excellent verbal and written communication skills). "
    "16. Required Technical Skills: Identify and list the technical skills required (e.g., Node.js, TypeScript, GraphQL). "
    "17. Required Soft Skills: Identify and list the soft skills required (e.g., Problem Solving, Excellent verbal and written communication skills). "
    "18. Keywords detected: Identify the top keywords in the job listing describing what skills and experiences the company is looking for in a candidate. Separate each key word as separate values for the key \"keywords_analysis\"."
)

    prompt = ChatPromptTemplate.from_messages([
        ("system", prompt_text),
        ("user", "{job_description_text}")
    ])

    logging.debug(f"Prompt created: {prompt}")

    chain = prompt | model

    input_data = {
        "job_description_text": job_description_text.strip()
    }
    logging.debug(f"Input data to chain.invoke: {input_data}")

    response = chain.invoke(input_data)
    logging.debug("Response received from chain.invoke")

    response_content = response.content.strip()
    logging.debug(f"Response content: {response_content}")

    if not response_content:
        logging.error("Received empty response from chain.invoke")
        raise ValueError("Received empty response from chain.invoke")

    try:
        if response_content.startswith("```json") and response_content.endswith("```"):
            response_content = response_content[7:-3].strip()
        
        response_json = json.loads(response_content)
        logging.debug(f"Response JSON: {response_json}")
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding JSON response: {str(e)}")
        logging.error(f"Response content was: {response_content}")
        
        response_content = response_content.replace("```json", "").replace("```", "").strip()
        try:
            response_json = json.loads(response_content)
            logging.debug(f"Recovered JSON: {response_json}")
        except json.JSONDecodeError as e:
            logging.error(f"Error decoding JSON response after recovery attempt: {str(e)}")
            raise ValueError(f"Error decoding JSON response: {str(e)}")

    required_keys = ["job_details", "company_information", "requirements_and_qualifications", "required_skill_sets", "Required_technical_skills", "Required_soft_skills", "keywords_analysis"]
    for key in required_keys:
        if key not in response_json:
            response_json[key] = None

    return response_json

def convert_to_date_format(date_str):
    if not date_str or date_str.lower() == "present":
        return date_str
    try:
        # Try parsing various formats
        for fmt in ("%B %Y", "%Y-%m-%d", "%Y-%m", "%m/%d/%Y", "%d/%m/%Y"):
            try:
                date = datetime.strptime(date_str, fmt)
                return date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        return date_str  # Return the original string if it cannot be parsed
    except Exception as e:
        logging.error(f"Error converting date {date_str}: {e}")
        return date_str


def get_resume_analysis(model, resume_text):
    prompt_text = """
    You are going to receive a resume from me. I need you to extract specific information and present it in a structured JSON format. Please provide the following details, ensuring each element is returned as a coherent string or a list of coherent strings, not individual letters:

    1. Header Text: Extract the main header or summary at the top of the resume. If it doesn't exist return NULL.
    2. Top Section Summary: Summarize the top section of the resume. If it doesn't exist return NULL.
    3. Top Section List of Achievements: List the achievements mentioned in the top section. List all of them.
    4. Education: Summarize the educational background including institutions attended, degrees received, and any mentioned GPA. If it doesn't exist return NULL.
    5. Bottom Section List of Achievements: List the achievements mentioned in the bottom section. If it doesn't exist return NULL.
    6. Achievements and Awards: List any certifications and awards mentioned, starting with the most recent. List them all as coherent strings in an array.
    7. Job Titles and Details: Extract details for up to six job titles including title, start date, end date, length, location, and description. Extract them in sequential order from the most recent being Job title 1.
       - Job Title 1: Extract the job title.
       - Job Title 1 Start Date: Extract the start date. Return it in 'YYYY-MM-DD' format.
       - Job Title 1 End Date: Extract the end date. Return it in 'YYYY-MM-DD' format or 'Present' if it is the current job.
       - Job Title 1 Length: Extract the job duration.
       - Job Title 1 Location: Extract the job location.
       - Job Title 1 Description: Extract the job description.
       - Repeat for Job Title 2 through Job Title 6.
       - If you don't find up to 6 job titles then return NULL for those job titles.
    8. Key Technical Skills: Identify and list technical skills. List them all as coherent strings in an array, in order of relevance.
    9. Key Soft Skills: Identify and list soft skills. List them all as coherent strings in an array, in order of relevance.
    10. Top Listed Skill Keyword: Identify the most frequently mentioned skill keyword.
    11. Second Most Top Listed Skill Keyword: Identify the second most frequently mentioned skill keyword.
    12. Third Most Top Listed Skill Keyword: Identify the third most frequently mentioned skill keyword.
    13. Fourth Most Top Listed Skill Keyword: Identify the fourth most frequently mentioned skill keyword.
    14. Most Recent Successful Project: Summarize the most recent successful project mentioned.
    15. Areas for Improvement: Suggest any areas where the resume could be improved or additional information that could be included.
    16. Questions About Experience: Note any unclear or concerning sections such as gaps or significant transitions.
    17. Resume Length: Count and return the number of characters in the resume.
    18. Top Challenge: Identify the most challenging project listed and provide its description.

    Here is the expected output JSON format. Ensure all values are returned as coherent strings or lists of coherent strings, and not as individual letters:

    {
      "header_text": "",
      "top_section_summary": "",
      "top_section_list_of_achievements": [],
      "education": "",
      "bottom_section_list_of_achievements": [],
      "achievements_and_awards": [],
      "job_title_1": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "job_title_2": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "job_title_3": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "job_title_4": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "job_title_5": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "job_title_6": {
        "title": "",
        "start_date": "",
        "end_date": "",
        "length": "",
        "location": "",
        "description": ""
      },
      "key_technical_skills": [],
      "key_soft_skills": [],
      "top_listed_skill_keyword": "",
      "second_most_top_listed_skill_keyword": "",
      "third_most_top_listed_skill_keyword": "",
      "fourth_most_top_listed_skill_keyword": "",
      "certifications_and_awards": [],
      "most_recent_successful_project": "",
      "areas_for_improvement": "",
      "questions_about_experience": "",
      "resume_length": "",
      "top_challenge": ""
    }

    Always return a value for every key in the correct format.
    """

    messages = [
        SystemMessage(content=prompt_text),
        HumanMessage(content=resume_text)
    ]

    logging.debug("Sending request to OpenAI with messages: %s", messages)
    response = model.invoke(messages)

    # Extract the content from the response
    if isinstance(response, dict) and 'choices' in response:
        response_content = response['choices'][0]['message']['content'].strip()
    elif hasattr(response, 'content'):
        response_content = response.content.strip()
    else:
        logging.error("Unexpected response format from OpenAI API: %s", response)
        raise ValueError("Unexpected response format from OpenAI API")

    # Remove triple backticks if they exist
    if response_content.startswith("```") and response_content.endswith("```"):
        response_content = response_content[3:-3].strip()

    logging.debug("Response Content: %s", response_content)

    try:
        response_json = json.loads(response_content)
    except json.JSONDecodeError as e:
        logging.error("JSON decode error: %s", e)
        logging.error("Response content that caused error: %s", response_content)
        raise ValueError("Failed to decode JSON response from OpenAI API")

    logging.debug("Response JSON: %s", response_json)

    # Convert date formats
    for i in range(1, 7):
        job_key = f"job_title_{i}"
        if job_key in response_json and isinstance(response_json[job_key], dict):
            for date_key in ['start_date', 'end_date']:
                if date_key in response_json[job_key]:
                    date_value = response_json[job_key][date_key]
                    if date_value.lower() == 'present':
                        response_json[job_key][date_key] = None  # Use None for 'Present'
                    else:
                        response_json[job_key][date_key] = convert_to_date_format(date_value)

    return response_json

def store_resume_analysis(user_id, response_json):
    with current_app.app_context():
        db_session = next(get_db())
        
        # Create a new Resume object
        resume = Resumes(
            user_id=user_id,
            header_text=response_json.get('header_text'),
            top_section_summary=response_json.get('top_section_summary'),
            top_section_list_of_achievements=json.dumps(response_json.get('top_section_list_of_achievements')),
            education=response_json.get('education'),
            bottom_section_list_of_achievements=json.dumps(response_json.get('bottom_section_list_of_achievements')),
            achievements_and_awards=json.dumps(response_json.get('achievements_and_awards')),
            job_title_1=json.dumps(response_json.get('job_title_1')),
            job_title_2=json.dumps(response_json.get('job_title_2')),
            job_title_3=json.dumps(response_json.get('job_title_3')),
            job_title_4=json.dumps(response_json.get('job_title_4')),
            job_title_5=json.dumps(response_json.get('job_title_5')),
            job_title_6=json.dumps(response_json.get('job_title_6')),
            key_technical_skills=json.dumps(response_json.get('key_technical_skills')),
            key_soft_skills=json.dumps(response_json.get('key_soft_skills')),
            top_listed_skill_keyword=response_json.get('top_listed_skill_keyword'),
            second_most_top_listed_skill_keyword=response_json.get('second_most_top_listed_skill_keyword'),
            third_most_top_listed_skill_keyword=response_json.get('third_most_top_listed_skill_keyword'),
            fourth_most_top_listed_skill_keyword=response_json.get('fourth_most_top_listed_skill_keyword'),
            most_recent_successful_project=response_json.get('most_recent_successful_project'),
            areas_for_improvement=response_json.get('areas_for_improvement'),
            questions_about_experience=response_json.get('questions_about_experience'),
            resume_length=response_json.get('resume_length'),
            top_challenge=response_json.get('top_challenge'),
            created_at=datetime.utcnow()
        )

        # Add the new resume data to the session and commit
        try:
            db_session.add(resume)
            db_session.commit()
            logging.info(f"Stored resume data for user_id: {user_id}")
        except Exception as e:
            logging.error(f"Error storing resume data: {str(e)}")
            db_session.rollback()

# Additional helper functions if needed

def cleanup_uploads_folder(app):
    save_dir = app.config['UPLOAD_FOLDER']
    files = glob.glob(os.path.join(save_dir, '*'))
    for file in files:
        try:
            os.remove(file)
            logging.info(f"Deleted file: {file}")
        except Exception as e:
            logging.error(f"Error deleting file {file}: {e}")

# functions for questions table
def add_question(db_session, question_data):
    new_question = Questions(**question_data)
    db_session.add(new_question)
    db_session.commit()
    return new_question

def get_questions(db_session, filters=None):
    query = db_session.query(Questions)
    if filters:
        query = query.filter_by(**filters)
    return query.all()

def update_question(db_session, question_id, update_data):
    question = db_session.query(Questions).filter_by(id=question_id).first()
    if question:
        for key, value in update_data.items():
            setattr(question, key, value)
        db_session.commit()
        return question
    return None

def delete_question(db_session, question_id):
    question = db_session.query(Questions).filter_by(id=question_id).first()
    if question:
        db_session.delete(question)
        db_session.commit()
        return True
    return False

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_user_by_username(username):
    session = SessionLocal()
    try:
        return session.query(User).filter_by(username=username).first()
    finally:
        session.close()

def extract_text_from_file(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_extension == '.docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif file_extension == '.pdf':
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    else:
        raise ValueError("Unsupported file format")

def store_job_description_data(response_json, user_id):
    with current_app.app_context():
        db_session = next(get_db())

        job_description = JobDescriptions(
            user_id=user_id,
            job_title=response_json['job_details'].get('title'),
            job_level=response_json['job_details'].get('level'),
            job_location=response_json['job_details'].get('location'),
            job_type=response_json['job_details'].get('type'),
            job_salary=response_json['job_details'].get('salary'),
            job_responsibilities=response_json['job_details'].get('responsibilities'),
            personal_qualifications=response_json['job_details'].get('personal_qualifications'),
            Required_technical_skills=response_json.get('Required_technical_skills'),
            Required_soft_skills=response_json.get('Required_soft_skills'),
            company_name=response_json['company_information'].get('name'),
            company_size=response_json['company_information'].get('size'),
            company_industry=response_json['company_information'].get('industry'),
            company_mission_and_values=response_json['company_information'].get('mission_and_values'),
            education_background=response_json['requirements_and_qualifications'].get('education_background'),
            required_professional_experiences=response_json['requirements_and_qualifications'].get('required_professional_experiences'),
            nice_to_have_experiences=response_json['requirements_and_qualifications'].get('nice_to_have_experiences'),
            required_skill_sets=response_json.get('required_skill_sets'),
            keywords_analysis=response_json.get('keywords_analysis'),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )

        try:
            db_session.add(job_description)
            db_session.commit()
            logging.info(f"Stored job description for user_id: {user_id}")
        except Exception as e:
            logging.error(f"Error storing job description: {str(e)}")
            db_session.rollback()
